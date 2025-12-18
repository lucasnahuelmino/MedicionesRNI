import os
from datetime import datetime
from io import BytesIO

import numpy as np
import pandas as pd
import streamlit as st
import plotly.express as px
import plotly.io as pio
from docx import Document
from docx.shared import Inches
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4

from utils.time_utils import calcular_tiempo_total_por_archivo, format_timedelta_long


def render_export_informes(df_localidad, df_filtrado_prov, localidad_seleccionada, titulo_scope):
    # ============================================================
    # 🖨️ EXPORTACIÓN DE INFORMES PDF / WORD
    # ============================================================

    with st.expander("🖨️ Generar informe PDF / Word", expanded=False):
        st.header("🖨️ Generar Informe con Gráficos y Datos Resumidos")

        if not st.session_state["tabla_maestra"].empty:
            # Usamos el mismo subset que se está viendo en pantalla:
            df_export = df_localidad.copy() if not df_localidad.empty else df_filtrado_prov.copy()

            # Si por algún motivo ese df está vacío, fallback a tabla completa
            if df_export.empty:
                df_export = st.session_state["tabla_maestra"].copy()

            # ========= ESTADÍSTICAS PARA EL RELATO =========
            df_export["Resultado"] = pd.to_numeric(df_export.get("Resultado", np.nan), errors="coerce")

            max_resultado = df_export["Resultado"].max() if "Resultado" in df_export.columns else None
            max_resultado_pct = None
            localidad_max = provincia_max = ccte_max = "N/D"
            fecha_hora_max = None

            if pd.notna(max_resultado):
                max_resultado_pct = max_resultado**2 / 3770 / 0.20021 * 100

                fila_max = df_export.loc[df_export["Resultado"].idxmax()]
                localidad_max = fila_max.get("Localidad", "N/D")
                provincia_max = fila_max.get("Provincia", "N/D")
                ccte_max = fila_max.get("CCTE", "N/D")

                # Fecha y hora del máximo
                if "FechaHora" in fila_max and pd.notna(fila_max["FechaHora"]):
                    fecha_hora_max = fila_max["FechaHora"]
                elif "Fecha" in fila_max and "Hora" in fila_max:
                    try:
                        fecha_hora_max = datetime.combine(fila_max["Fecha"], fila_max["Hora"])
                    except Exception:
                        fecha_hora_max = fila_max.get("Fecha", None)
                else:
                    fecha_hora_max = fila_max.get("Fecha", None)

            # Rango de fechas trabajadas
            fecha_min = fecha_max_med = None
            if "Fecha" in df_export.columns:
                fechas = pd.to_datetime(df_export["Fecha"], dayfirst=True, errors="coerce")
                if fechas.notna().any():
                    fecha_min = fechas.min().date()
                    fecha_max_med = fechas.max().date()

            # Tiempo total trabajado (según Nombre Archivo + Fecha/Hora)
            tiempo_total_trabajado = calcular_tiempo_total_por_archivo(df_export)

            # Sondas utilizadas
            sondas_uniq = []
            if "Sonda" in df_export.columns:
                sondas_uniq = sorted(df_export["Sonda"].dropna().astype(str).unique().tolist())

            # ========= DESGLOSE POR MES =========
            resumen_mensual_export = pd.DataFrame()
            if "FechaHora" in df_export.columns and df_export["FechaHora"].notna().any():
                df_tmp = df_export.copy()
                df_tmp["FechaHora"] = pd.to_datetime(df_tmp["FechaHora"], errors="coerce")
                df_tmp["Mes"] = df_tmp["FechaHora"].dt.to_period("M")

                if df_tmp["Mes"].notna().any():
                    resumen_mensual_export = df_tmp.groupby("Mes").agg(
                        Cantidad_puntos=("Resultado", "count"),
                        Fecha_inicio=("FechaHora", "min"),
                        Fecha_fin=("FechaHora", "max"),
                        Localidades_trabajadas=("Localidad", lambda x: ", ".join(sorted(x.dropna().unique())))
                    ).reset_index()

                    # Tiempo por mes
                    def _tiempo_mes(g_mes):
                        return format_timedelta_long(calcular_tiempo_total_por_archivo(g_mes))

                    filas_tiempo_export = []
                    for mes, g_mes in df_tmp.groupby("Mes"):
                        filas_tiempo_export.append({
                            "Mes": mes,
                            "Horas_trabajadas": _tiempo_mes(g_mes)
                        })
                    tiempos_mes = pd.DataFrame(filas_tiempo_export)

                    resumen_mensual_export = resumen_mensual_export.merge(tiempos_mes, on="Mes")

            # ========= TABLA DE EXPEDIENTES =========
            expedientes_df = pd.DataFrame()
            if "Expediente" in df_export.columns:
                expedientes_df = df_export.groupby("Expediente").agg(
                    Cantidad_puntos=("Resultado", "count"),
                    CCTE=("CCTE", lambda x: ", ".join(sorted(x.dropna().unique()))),
                    Provincias=("Provincia", lambda x: ", ".join(sorted(x.dropna().unique()))),
                    Localidades=("Localidad", lambda x: ", ".join(sorted(x.dropna().unique()))),
                    Max_Vm=("Resultado", "max")
                ).reset_index()
                expedientes_df = expedientes_df.sort_values(by="Max_Vm", ascending=False)

            # ========= GRÁFICO SOLO DEL ÁMBITO ACTUAL =========
            df_graf_export = df_export.copy()
            if {"Provincia", "CCTE", "Localidad"}.issubset(df_graf_export.columns):
                resumen_export = (
                    df_graf_export
                    .groupby(["Provincia", "CCTE"])["Localidad"]
                    .nunique()
                    .reset_index(name="CantidadLocalidades")
                )
            else:
                resumen_export = pd.DataFrame()

            fig_bar_export = None
            if not resumen_export.empty:
                fig_bar_export = px.bar(
                    resumen_export,
                    x="Provincia",
                    y="CantidadLocalidades",
                    color="CCTE",
                    text="CantidadLocalidades",
                    barmode="group",
                    title="Localidades por Provincia y CCTE (ámbito del informe)",
                    color_discrete_sequence=px.colors.qualitative.Set2
                )
                fig_bar_export.update_layout(template="plotly_white")

            # ========= OPCIONES DE EXPORTACIÓN =========
            col_exp1, _ = st.columns(2)
            formato = col_exp1.radio("Formato de exportación", ["Word (.docx)", "PDF (.pdf)"], horizontal=True)

            if st.button("📄 Generar Informe"):
                localidad_nombre = localidad_seleccionada or "General"
                fecha_str = datetime.now().strftime("%Y%m%d_%H%M")

                # ============================================================
                # 🧾 WORD (sin header azul, con logo y tablas)
                # ============================================================
                if formato == "Word (.docx)":
                    doc = Document()

                    # Logo arriba del informe (sin header azul)
                    if os.path.exists("assets/enacom_logo.png"):
                        doc.add_picture("assets/enacom_logo.png", width=Inches(2.5))

                    doc.add_heading(f"Informe de Mediciones RNI - {localidad_nombre}", level=1)
                    doc.add_paragraph(f"Ámbito del informe: {titulo_scope}")
                    doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
                    doc.add_paragraph(f"Total de puntos medidos: {len(df_export)}")

                    if pd.notna(max_resultado):
                        p_max = doc.add_paragraph()
                        p_max.add_run("Resultado máximo registrado: ").bold = True
                        p_max.add_run(f"{max_resultado:.2f} V/m")
                        if max_resultado_pct is not None:
                            p_max.add_run(f" ({max_resultado_pct:.2f} % del límite)")

                        p_ub = doc.add_paragraph()
                        p_ub.add_run("Ubicación del máximo: ").bold = True
                        p_ub.add_run(f"{localidad_max}, {provincia_max} (CCTE {ccte_max})")

                        if fecha_hora_max is not None:
                            p_fm = doc.add_paragraph()
                            p_fm.add_run("Fecha y hora del máximo: ").bold = True
                            p_fm.add_run(str(fecha_hora_max))

                    if fecha_min and fecha_max_med:
                        p_f = doc.add_paragraph()
                        p_f.add_run("Rango de fechas de medición: ").bold = True
                        p_f.add_run(f"{fecha_min.strftime('%d/%m/%Y')} a {fecha_max_med.strftime('%d/%m/%Y')}")

                    if tiempo_total_trabajado.total_seconds() > 0:
                        p_t = doc.add_paragraph()
                        p_t.add_run("Tiempo total estimado de medición: ").bold = True
                        p_t.add_run(format_timedelta_long(tiempo_total_trabajado))

                    if sondas_uniq:
                        p_s = doc.add_paragraph()
                        p_s.add_run("Sondas utilizadas: ").bold = True
                        p_s.add_run(", ".join(sondas_uniq))

                    doc.add_paragraph(" ")

                    # --- Gráfico principal (ámbito actual) ---
                    if fig_bar_export is not None:
                        img_bytes = BytesIO()
                        pio.write_image(fig_bar_export, img_bytes, format="png")
                        img_bytes.seek(0)
                        doc.add_picture(img_bytes, width=Inches(5.5))
                        doc.add_paragraph("Gráfico de Localidades por Provincia y CCTE (ámbito del informe).")

                    # --- Desglose por mes (tabla) ---
                    if not resumen_mensual_export.empty:
                        doc.add_heading("Desglose por mes", level=2)
                        table = doc.add_table(rows=1, cols=5)
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = "Mes"
                        hdr_cells[1].text = "Puntos"
                        hdr_cells[2].text = "Horas trabajadas"
                        hdr_cells[3].text = "Fecha inicio"
                        hdr_cells[4].text = "Fecha fin"

                        for _, row in resumen_mensual_export.iterrows():
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(row["Mes"])
                            row_cells[1].text = str(row["Cantidad_puntos"])
                            row_cells[2].text = row["Horas_trabajadas"]
                            fi = row["Fecha_inicio"]
                            ff = row["Fecha_fin"]
                            row_cells[3].text = fi.strftime("%d/%m/%Y %H:%M") if pd.notna(fi) else "-"
                            row_cells[4].text = ff.strftime("%d/%m/%Y %H:%M") if pd.notna(ff) else "-"

                    # --- Tabla de expedientes ---
                    if not expedientes_df.empty:
                        doc.add_heading("Resumen por expediente", level=2)
                        table_e = doc.add_table(rows=1, cols=6)
                        hdr = table_e.rows[0].cells
                        hdr[0].text = "Expediente"
                        hdr[1].text = "Puntos"
                        hdr[2].text = "Max (V/m)"
                        hdr[3].text = "CCTE"
                        hdr[4].text = "Provincias"
                        hdr[5].text = "Localidades"

                        for _, row in expedientes_df.iterrows():
                            r = table_e.add_row().cells
                            r[0].text = str(row["Expediente"])
                            r[1].text = str(row["Cantidad_puntos"])
                            r[2].text = f"{row['Max_Vm']:.2f}" if pd.notna(row["Max_Vm"]) else "-"
                            r[3].text = str(row["CCTE"])
                            r[4].text = str(row["Provincias"])
                            r[5].text = str(row["Localidades"])

                    ruta_doc = f"Informe_RNI_{localidad_nombre}_{fecha_str}.docx"
                    doc.save(ruta_doc)
                    with open(ruta_doc, "rb") as f:
                        st.download_button(
                            label="⬇️ Descargar Informe Word",
                            data=f,
                            file_name=ruta_doc,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                # ============================================================
                # 📘 PDF (sin header azul, con desglose)
                # ============================================================
                else:
                    ruta_pdf = f"Informe_RNI_{localidad_nombre}_{fecha_str}.pdf"
                    buffer = BytesIO()
                    pdf = SimpleDocTemplate(buffer, pagesize=A4)
                    styles = getSampleStyleSheet()
                    style_title = styles["Title"]
                    style_sub = styles["Heading2"]
                    style_normal = styles["Normal"]

                    story = []

                    # Logo si está disponible
                    if os.path.exists("assets/enacom_logo.png"):
                        story.append(RLImage("assets/enacom_logo.png", width=200, height=60))
                        story.append(Spacer(1, 12))

                    story.append(Paragraph("Informe de Mediciones RNI", style_title))
                    story.append(Spacer(1, 6))
                    story.append(Paragraph(f"Ámbito del informe: {titulo_scope}", style_sub))
                    story.append(Spacer(1, 12))

                    story.append(Paragraph(f"<b>Localidad seleccionada:</b> {localidad_nombre}", style_normal))
                    story.append(Paragraph(f"<b>Fecha de generación:</b> {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", style_normal))
                    story.append(Paragraph(f"<b>Total de puntos medidos:</b> {len(df_export)}", style_normal))

                    if pd.notna(max_resultado):
                        story.append(Paragraph(
                            f"<b>Resultado máximo registrado:</b> {max_resultado:.2f} V/m"
                            + (f" ({max_resultado_pct:.2f} % del límite)" if max_resultado_pct is not None else ""),
                            style_normal
                        ))
                        story.append(Paragraph(
                            f"<b>Ubicación del máximo:</b> {localidad_max}, {provincia_max} (CCTE {ccte_max})",
                            style_normal
                        ))
                        if fecha_hora_max is not None:
                            story.append(Paragraph(
                                f"<b>Fecha y hora del máximo:</b> {fecha_hora_max}",
                                style_normal
                            ))

                    if fecha_min and fecha_max_med:
                        story.append(Paragraph(
                            f"<b>Rango de fechas de medición:</b> {fecha_min.strftime('%d/%m/%Y')} a {fecha_max_med.strftime('%d/%m/%Y')}",
                            style_normal
                        ))

                    if tiempo_total_trabajado.total_seconds() > 0:
                        story.append(Paragraph(
                            f"<b>Tiempo total estimado de medición:</b> {format_timedelta_long(tiempo_total_trabajado)}",
                            style_normal
                        ))

                    if sondas_uniq:
                        story.append(Paragraph(
                            f"<b>Sondas utilizadas:</b> {', '.join(sondas_uniq)}",
                            style_normal
                        ))

                    story.append(Spacer(1, 16))

                    # Gráfico (si hay)
                    if fig_bar_export is not None:
                        img_bytes = BytesIO()
                        pio.write_image(fig_bar_export, img_bytes, format="png")
                        img_bytes.seek(0)
                        story.append(RLImage(img_bytes, width=400, height=250))
                        story.append(Paragraph("Gráfico de Localidades por Provincia y CCTE (ámbito del informe)", styles["Italic"]))
                        story.append(Spacer(1, 16))

                    # Desglose mensual (en texto)
                    if not resumen_mensual_export.empty:
                        story.append(Paragraph("<b>Desglose por mes</b>", style_sub))
                        story.append(Spacer(1, 6))
                        for _, row in resumen_mensual_export.iterrows():
                            fi = row["Fecha_inicio"]
                            ff = row["Fecha_fin"]
                            texto = (
                                f"Mes {row['Mes']}: {row['Cantidad_puntos']} puntos, "
                                f"horas trabajadas: {row['Horas_trabajadas']}, "
                                f"localidades: {row['Localidades_trabajadas']}. "
                            )
                            if pd.notna(fi) and pd.notna(ff):
                                texto += f"({fi.strftime('%d/%m/%Y %H:%M')} a {ff.strftime('%d/%m/%Y %H:%M')})"
                            story.append(Paragraph(texto, style_normal))
                        story.append(Spacer(1, 12))

                    # Tabla de expedientes (en texto)
                    if not expedientes_df.empty:
                        story.append(Paragraph("<b>Resumen por expediente</b>", style_sub))
                        story.append(Spacer(1, 6))
                        for _, row in expedientes_df.iterrows():
                            texto = (
                                f"Expediente {row['Expediente']}: "
                                f"{row['Cantidad_puntos']} puntos, "
                                f"máx {row['Max_Vm']:.2f} V/m, "
                                f"CCTE: {row['CCTE']}, "
                                f"Provincias: {row['Provincias']}, "
                                f"Localidades: {row['Localidades']}."
                            )
                            story.append(Paragraph(texto, style_normal))
                        story.append(Spacer(1, 12))

                    pdf.build(story)
                    buffer.seek(0)
                    st.download_button(
                        label="⬇️ Descargar Informe PDF",
                        data=buffer,
                        file_name=ruta_pdf,
                        mime="application/pdf"
                    )
